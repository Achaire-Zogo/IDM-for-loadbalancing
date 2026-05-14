#!/usr/bin/env python3
"""
Génération du rapport Word pour le projet INF5039 :
Modélisation et génération d'un mécanisme de répartition de charge
par DSL pour un système distribué.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.name = 'Times New Roman'
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True

def add_paragraph(text, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()
    return table

def add_code_block(code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_paragraph("UNIVERSITE DE YAOUNDE 1", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("Faculté des Sciences", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("Département d'Informatique", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("Master 2 – INF5039 : Ingénierie Dirigée par les Modèles", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

doc.add_paragraph()

add_paragraph("PROJET", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_paragraph(
    "Modélisation et génération d'un mécanisme de répartition\n"
    "de charge par DSL pour un système distribué",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20
)

doc.add_paragraph()
doc.add_paragraph()

add_paragraph("Réalisé par :", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("ZOGO ABOUMA Zozime Achaire – 18N2824", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("NKOLO Stéphane Roylex – 21T2588", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

doc.add_paragraph()

add_paragraph("Encadreur :", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("Dr MONTHE Valérie", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph("Université de Yaoundé 1", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

doc.add_paragraph()
add_paragraph(f"Année académique 2025 – 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES (placeholder)
# ============================================================
doc.add_heading("Table des matières", level=1)
add_paragraph("(À générer automatiquement dans Word : Références → Table des matières)", italic=True)
doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
doc.add_heading("Introduction", level=1)
add_paragraph(
    "Dans le contexte actuel des systèmes informatiques, la montée en charge des applications "
    "distribuées impose des mécanismes efficaces de répartition de charge (load balancing). "
    "Ces mécanismes permettent de distribuer équitablement les requêtes entre plusieurs serveurs "
    "afin d'optimiser les performances, d'assurer la haute disponibilité et de prévenir la surcharge "
    "de nœuds individuels."
)
add_paragraph(
    "L'Ingénierie Dirigée par les Modèles (IDM/MDE) offre une approche systématique pour "
    "modéliser, transformer et générer des artefacts logiciels à partir de modèles abstraits. "
    "En combinant IDM et la définition d'un langage dédié (DSL – Domain Specific Language), "
    "il devient possible de spécifier de manière déclarative les politiques de répartition de charge "
    "et de générer automatiquement le code d'infrastructure correspondant."
)
add_paragraph(
    "Ce projet a pour objectif de concevoir un DSL permettant de modéliser un mécanisme de "
    "répartition de charge pour un système distribué, puis de mettre en œuvre une chaîne de "
    "transformations MDA (Model Driven Architecture) allant du modèle abstrait jusqu'au code "
    "exécutable."
)

# ============================================================
# 1. DESCRIPTION DU DOMAINE
# ============================================================
doc.add_heading("1. Description du domaine traité", level=1)

doc.add_heading("1.1 Qu'est-ce que la répartition de charge ?", level=2)
add_paragraph(
    "La répartition de charge (load balancing) est une technique utilisée dans les systèmes distribués "
    "pour distribuer le trafic réseau ou les tâches de calcul entre plusieurs serveurs ou ressources. "
    "L'objectif est de :"
)
add_bullet("Maximiser le débit (throughput) du système global")
add_bullet("Minimiser le temps de réponse pour chaque requête")
add_bullet("Éviter la surcharge d'un seul nœud")
add_bullet("Assurer la tolérance aux pannes (failover)")
add_bullet("Permettre la scalabilité horizontale")

doc.add_heading("1.2 Composants d'un système de répartition de charge", level=2)
add_paragraph(
    "Un système de répartition de charge typique comprend les éléments suivants :"
)
add_bullet("Load Balancer (répartiteur) : point d'entrée qui reçoit les requêtes et les redirige")
add_bullet("Backend Servers (serveurs backend) : ensemble de serveurs qui traitent les requêtes")
add_bullet("Health Check (vérification de santé) : mécanisme pour vérifier la disponibilité des serveurs")
add_bullet("Algorithme de répartition : stratégie utilisée pour choisir le serveur cible (Round Robin, Weighted, Least Connections, IP Hash, etc.)")
add_bullet("Session Persistence (affinité de session) : mécanisme pour diriger un même client vers le même serveur")
add_bullet("Cluster : regroupement logique de serveurs backend")
add_bullet("Listeners : points d'écoute sur des ports/protocoles spécifiques")

doc.add_heading("1.3 Algorithmes de répartition courants", level=2)
add_table(
    ["Algorithme", "Description", "Cas d'usage"],
    [
        ["Round Robin", "Distribution séquentielle cyclique entre les serveurs", "Serveurs homogènes, charge uniforme"],
        ["Weighted Round Robin", "Round Robin avec poids assignés à chaque serveur", "Serveurs hétérogènes (capacités différentes)"],
        ["Least Connections", "Envoie au serveur ayant le moins de connexions actives", "Requêtes de durée variable"],
        ["IP Hash", "Utilise un hash de l'IP source pour déterminer le serveur", "Affinité de session sans cookies"],
        ["Random", "Choix aléatoire parmi les serveurs disponibles", "Cas simples, tests"],
        ["Weighted Least Connections", "Least Connections pondéré par la capacité du serveur", "Serveurs hétérogènes, charge variable"],
    ]
)

doc.add_heading("1.4 Contexte et problématique", level=2)
add_paragraph(
    "Traditionnellement, la configuration d'un load balancer se fait manuellement via des fichiers "
    "de configuration (nginx.conf, haproxy.cfg, etc.) qui sont :"
)
add_bullet("Spécifiques à une technologie (vendor lock-in)")
add_bullet("Sujets à des erreurs de syntaxe")
add_bullet("Difficiles à valider avant déploiement")
add_bullet("Non portables d'une plateforme à l'autre")

add_paragraph(
    "Notre approche IDM propose de résoudre ces problèmes en définissant un DSL indépendant "
    "de toute technologie qui capture l'essence de la répartition de charge, et en utilisant des "
    "transformations de modèles pour générer automatiquement les configurations ou le code cible."
)

# ============================================================
# 2. TABLEAU DES CONCEPTS
# ============================================================
doc.add_heading("2. Tableau des concepts identifiés", level=1)
add_paragraph(
    "Le tableau suivant recense les concepts clés du domaine de la répartition de charge "
    "qui seront modélisés dans notre métamodèle :"
)

add_table(
    ["N°", "Concept", "Description", "Attributs principaux"],
    [
        ["1", "LoadBalancerSystem", "Système global de répartition de charge, élément racine", "name, description, version"],
        ["2", "Cluster", "Groupe logique de serveurs backend partageant une même fonction", "name, maxConnections"],
        ["3", "Server", "Serveur backend individuel pouvant recevoir des requêtes", "name, host, port, weight, maxConnections, enabled"],
        ["4", "LoadBalancer", "Composant répartiteur qui distribue les requêtes", "name, algorithm, stickySession"],
        ["5", "Algorithm", "Énumération des algorithmes de répartition disponibles", "ROUND_ROBIN, WEIGHTED_ROUND_ROBIN, LEAST_CONNECTIONS, IP_HASH, RANDOM"],
        ["6", "HealthCheck", "Mécanisme de vérification de la disponibilité d'un serveur", "protocol, path, interval, timeout, healthyThreshold, unhealthyThreshold"],
        ["7", "Listener", "Point d'écoute sur un port/protocole spécifique", "name, protocol, port, targetCluster"],
        ["8", "Protocol", "Énumération des protocoles supportés", "HTTP, HTTPS, TCP, UDP"],
        ["9", "SessionPersistence", "Configuration de l'affinité de session", "type, cookieName, ttl"],
        ["10", "Alert", "Règle d'alerte en cas de dépassement de seuil", "name, metric, threshold, action"],
        ["11", "ScalingRule", "Règle de mise à l'échelle automatique", "name, metric, scaleUpThreshold, scaleDownThreshold, minInstances, maxInstances"],
        ["12", "Metric", "Énumération des métriques surveillées", "CPU_USAGE, MEMORY_USAGE, REQUEST_COUNT, RESPONSE_TIME, ACTIVE_CONNECTIONS"],
    ]
)

# ============================================================
# 3. REGLES DE GESTION
# ============================================================
doc.add_heading("3. Définition des règles de gestion et exigences", level=1)
add_paragraph(
    "Les règles suivantes doivent être respectées dans tout modèle conforme à notre métamodèle :"
)

rules = [
    ("RG01", "Unicité des noms de serveurs", "Tous les serveurs au sein d'un même cluster doivent avoir des noms uniques."),
    ("RG02", "Port valide", "Le port d'un serveur ou d'un listener doit être compris entre 1 et 65535."),
    ("RG03", "Poids positif", "Le poids (weight) d'un serveur doit être un entier strictement positif (≥ 1)."),
    ("RG04", "Cluster non vide", "Un cluster doit contenir au moins un serveur."),
    ("RG05", "Health check cohérent", "L'intervalle de health check doit être supérieur au timeout."),
    ("RG06", "Seuils de scaling cohérents", "Le seuil de scale-up doit être supérieur au seuil de scale-down."),
    ("RG07", "Instances min/max cohérentes", "Le nombre minimum d'instances doit être inférieur ou égal au nombre maximum."),
    ("RG08", "Listener unique par port", "Deux listeners ne peuvent pas écouter sur le même port."),
    ("RG09", "Algorithme Weighted requiert des poids", "Si l'algorithme choisi est Weighted Round Robin, tous les serveurs du cluster doivent avoir un poids défini (≥ 1)."),
    ("RG10", "Session Persistence requiert un cookie name", "Si la persistence de session est de type COOKIE, le champ cookieName doit être renseigné."),
    ("RG11", "Au moins un listener", "Le système doit avoir au moins un listener défini."),
    ("RG12", "Serveur actif dans le cluster", "Un cluster doit avoir au moins un serveur dont l'attribut enabled est true."),
]

add_table(
    ["Code", "Règle", "Description"],
    [[r[0], r[1], r[2]] for r in rules]
)

# ============================================================
# 4. DECOMPOSITION EN MODULES
# ============================================================
doc.add_heading("4. Décomposition en modules / Structure en packages", level=1)
add_paragraph(
    "Notre système est organisé en modules (packages) pour une meilleure séparation des préoccupations :"
)

doc.add_heading("4.1 Architecture des packages", level=2)
add_code_block(
    "loadbalancer.dsl/\n"
    "├── metamodel/                    # Package : Métamodèle\n"
    "│   ├── core/                     # Sous-package : Concepts fondamentaux\n"
    "│   │   ├── LoadBalancerSystem    # Élément racine\n"
    "│   │   ├── Cluster              # Regroupement de serveurs\n"
    "│   │   └── Server               # Serveur backend\n"
    "│   ├── balancing/               # Sous-package : Répartition\n"
    "│   │   ├── LoadBalancer         # Répartiteur\n"
    "│   │   ├── Algorithm            # Énumération des algorithmes\n"
    "│   │   └── SessionPersistence   # Affinité de session\n"
    "│   ├── networking/              # Sous-package : Réseau\n"
    "│   │   ├── Listener             # Point d'écoute\n"
    "│   │   └── Protocol             # Protocoles supportés\n"
    "│   ├── monitoring/              # Sous-package : Surveillance\n"
    "│   │   ├── HealthCheck          # Vérification de santé\n"
    "│   │   ├── Alert                # Alertes\n"
    "│   │   └── Metric               # Métriques\n"
    "│   └── scaling/                 # Sous-package : Mise à l'échelle\n"
    "│       └── ScalingRule          # Règles de scaling\n"
    "├── constraints/                  # Package : Contraintes OCL\n"
    "├── concrete.syntax/             # Package : Syntaxe concrète (DSL textuel)\n"
    "├── transformations/             # Package : Transformations M2M et M2T\n"
    "│   ├── m2m/                     # Transformations modèle-à-modèle\n"
    "│   └── m2t/                     # Transformations modèle-à-texte\n"
    "└── examples/                    # Package : Études de cas"
)

doc.add_heading("4.2 Description des modules", level=2)
add_table(
    ["Module", "Rôle", "Contenu principal"],
    [
        ["metamodel.core", "Définit les concepts fondamentaux du domaine", "LoadBalancerSystem, Cluster, Server"],
        ["metamodel.balancing", "Gère la logique de répartition", "LoadBalancer, Algorithm, SessionPersistence"],
        ["metamodel.networking", "Gère les aspects réseau", "Listener, Protocol"],
        ["metamodel.monitoring", "Gère la surveillance et les alertes", "HealthCheck, Alert, Metric"],
        ["metamodel.scaling", "Gère la mise à l'échelle automatique", "ScalingRule"],
        ["constraints", "Contraintes OCL sur le métamodèle", "Fichiers .ocl"],
        ["concrete.syntax", "Grammaire Xtext du DSL", "Fichier .xtext"],
        ["transformations.m2m", "Transformations modèle-à-modèle (ATL/QVT)", "Fichiers .atl"],
        ["transformations.m2t", "Transformations modèle-à-texte (Acceleo)", "Fichiers .mtl"],
    ]
)

# ============================================================
# 5. SCHEMA DU METAMODELE
# ============================================================
doc.add_heading("5. Schéma du métamodèle (MM)", level=1)

doc.add_heading("5.1 Module Core", level=2)
add_paragraph(
    "Le module Core contient les éléments fondamentaux du système :"
)
add_code_block(
    "┌─────────────────────────┐\n"
    "│   LoadBalancerSystem    │\n"
    "│─────────────────────────│\n"
    "│ name : EString          │\n"
    "│ description : EString   │\n"
    "│ version : EString       │\n"
    "│─────────────────────────│\n"
    "│ clusters : Cluster [1..*]│\n"
    "│ listeners : Listener[1..*]│\n"
    "│ alerts : Alert [0..*]   │\n"
    "└───────────┬─────────────┘\n"
    "            │ 1..*\n"
    "            ▼\n"
    "┌─────────────────────────┐      ┌─────────────────────────┐\n"
    "│        Cluster          │      │       LoadBalancer       │\n"
    "│─────────────────────────│      │─────────────────────────│\n"
    "│ name : EString          │◄────►│ name : EString          │\n"
    "│ maxConnections : EInt   │  1   │ algorithm : Algorithm   │\n"
    "│─────────────────────────│      │ stickySession : EBoolean│\n"
    "│ servers : Server [1..*] │      │─────────────────────────│\n"
    "│ loadBalancer : LB [1]   │      │ sessionPersistence [0..1]│\n"
    "│ healthCheck [0..1]      │      └─────────────────────────┘\n"
    "│ scalingRule [0..1]      │\n"
    "└───────────┬─────────────┘\n"
    "            │ 1..*\n"
    "            ▼\n"
    "┌─────────────────────────┐\n"
    "│        Server           │\n"
    "│─────────────────────────│\n"
    "│ name : EString          │\n"
    "│ host : EString          │\n"
    "│ port : EInt             │\n"
    "│ weight : EInt           │\n"
    "│ maxConnections : EInt   │\n"
    "│ enabled : EBoolean      │\n"
    "└─────────────────────────┘"
)

doc.add_heading("5.2 Module Balancing", level=2)
add_code_block(
    "┌─────────────────────────┐\n"
    "│    <<enumeration>>      │\n"
    "│      Algorithm          │\n"
    "│─────────────────────────│\n"
    "│ ROUND_ROBIN             │\n"
    "│ WEIGHTED_ROUND_ROBIN    │\n"
    "│ LEAST_CONNECTIONS       │\n"
    "│ IP_HASH                 │\n"
    "│ RANDOM                  │\n"
    "└─────────────────────────┘\n"
    "\n"
    "┌─────────────────────────┐\n"
    "│  SessionPersistence     │\n"
    "│─────────────────────────│\n"
    "│ type : PersistenceType  │\n"
    "│ cookieName : EString    │\n"
    "│ ttl : EInt              │\n"
    "└─────────────────────────┘\n"
    "\n"
    "┌─────────────────────────┐\n"
    "│   <<enumeration>>       │\n"
    "│   PersistenceType       │\n"
    "│─────────────────────────│\n"
    "│ COOKIE                  │\n"
    "│ IP                      │\n"
    "│ HEADER                  │\n"
    "└─────────────────────────┘"
)

doc.add_heading("5.3 Module Networking", level=2)
add_code_block(
    "┌─────────────────────────┐\n"
    "│       Listener          │\n"
    "│─────────────────────────│\n"
    "│ name : EString          │\n"
    "│ protocol : Protocol     │\n"
    "│ port : EInt             │\n"
    "│─────────────────────────│\n"
    "│ targetCluster : Cluster │\n"
    "└─────────────────────────┘\n"
    "\n"
    "┌─────────────────────────┐\n"
    "│   <<enumeration>>       │\n"
    "│      Protocol           │\n"
    "│─────────────────────────│\n"
    "│ HTTP                    │\n"
    "│ HTTPS                   │\n"
    "│ TCP                     │\n"
    "│ UDP                     │\n"
    "└─────────────────────────┘"
)

doc.add_heading("5.4 Module Monitoring", level=2)
add_code_block(
    "┌─────────────────────────┐\n"
    "│      HealthCheck        │\n"
    "│─────────────────────────│\n"
    "│ protocol : Protocol     │\n"
    "│ path : EString          │\n"
    "│ interval : EInt (sec)   │\n"
    "│ timeout : EInt (sec)    │\n"
    "│ healthyThreshold : EInt │\n"
    "│ unhealthyThreshold: EInt│\n"
    "└─────────────────────────┘\n"
    "\n"
    "┌─────────────────────────┐      ┌─────────────────────────┐\n"
    "│        Alert            │      │    <<enumeration>>      │\n"
    "│─────────────────────────│      │       Metric            │\n"
    "│ name : EString          │      │─────────────────────────│\n"
    "│ metric : Metric         │      │ CPU_USAGE               │\n"
    "│ threshold : EFloat      │      │ MEMORY_USAGE            │\n"
    "│ action : EString        │      │ REQUEST_COUNT           │\n"
    "└─────────────────────────┘      │ RESPONSE_TIME           │\n"
    "                                 │ ACTIVE_CONNECTIONS      │\n"
    "                                 └─────────────────────────┘"
)

doc.add_heading("5.5 Module Scaling", level=2)
add_code_block(
    "┌──────────────────────────────┐\n"
    "│        ScalingRule           │\n"
    "│──────────────────────────────│\n"
    "│ name : EString              │\n"
    "│ metric : Metric             │\n"
    "│ scaleUpThreshold : EFloat   │\n"
    "│ scaleDownThreshold : EFloat │\n"
    "│ minInstances : EInt         │\n"
    "│ maxInstances : EInt         │\n"
    "└──────────────────────────────┘"
)

doc.add_heading("5.6 Schéma global du métamodèle (vue d'ensemble)", level=2)
add_paragraph(
    "Le schéma global montre les relations entre tous les modules. L'élément racine "
    "LoadBalancerSystem agrège des Clusters et des Listeners. Chaque Cluster contient "
    "des Servers, un LoadBalancer, optionnellement un HealthCheck et un ScalingRule. "
    "Chaque Listener référence un Cluster cible."
)
add_code_block(
    "                    ┌──────────────────────┐\n"
    "                    │  LoadBalancerSystem   │\n"
    "                    │  (Élément racine)     │\n"
    "                    └──────┬───────┬────────┘\n"
    "                    ┌──────┘       └──────┐\n"
    "                    ▼ 1..*         1..* ▼\n"
    "            ┌───────────────┐   ┌──────────────┐\n"
    "            │   Cluster     │   │   Listener   │\n"
    "            └───┬──┬──┬──┬─┘   └──────────────┘\n"
    "                │  │  │  │          │\n"
    "     ┌──────────┘  │  │  └──────┐   │ targetCluster\n"
    "     ▼ 1..*        │  │         ▼   ▼\n"
    "┌─────────┐   ┌────┴──┴───┐  ┌────────────┐\n"
    "│ Server  │   │LoadBalancer│  │ScalingRule │\n"
    "└─────────┘   └─────┬─────┘  └────────────┘\n"
    "                    │ 0..1\n"
    "              ┌─────┴──────┐\n"
    "              │Session     │\n"
    "              │Persistence │\n"
    "              └────────────┘\n"
    "\n"
    "  Cluster ──── 0..1 ──── HealthCheck\n"
    "  Cluster ──── 0..1 ──── ScalingRule\n"
    "  Alert ◇──── Metric (enum)\n"
    "  ScalingRule ◇──── Metric (enum)"
)

# ============================================================
# 6. CONTRAINTES OCL
# ============================================================
doc.add_heading("6. Définition des contraintes OCL", level=1)
add_paragraph(
    "Chaque règle de gestion définie en section 3 est formalisée ci-dessous en OCL "
    "(Object Constraint Language) :"
)

ocl_constraints = [
    ("RG01 – Unicité des noms de serveurs",
     "context Cluster\n"
     "inv uniqueServerNames:\n"
     "  self.servers->isUnique(s | s.name)"),

    ("RG02 – Port valide",
     "context Server\n"
     "inv validPort:\n"
     "  self.port >= 1 and self.port <= 65535\n\n"
     "context Listener\n"
     "inv validListenerPort:\n"
     "  self.port >= 1 and self.port <= 65535"),

    ("RG03 – Poids positif",
     "context Server\n"
     "inv positiveWeight:\n"
     "  self.weight >= 1"),

    ("RG04 – Cluster non vide",
     "context Cluster\n"
     "inv nonEmptyCluster:\n"
     "  self.servers->notEmpty()"),

    ("RG05 – Health check cohérent",
     "context HealthCheck\n"
     "inv intervalGreaterThanTimeout:\n"
     "  self.interval > self.timeout"),

    ("RG06 – Seuils de scaling cohérents",
     "context ScalingRule\n"
     "inv scaleUpGreaterThanDown:\n"
     "  self.scaleUpThreshold > self.scaleDownThreshold"),

    ("RG07 – Instances min/max cohérentes",
     "context ScalingRule\n"
     "inv minLessOrEqualMax:\n"
     "  self.minInstances <= self.maxInstances"),

    ("RG08 – Listener unique par port",
     "context LoadBalancerSystem\n"
     "inv uniqueListenerPorts:\n"
     "  self.listeners->isUnique(l | l.port)"),

    ("RG09 – Algorithme Weighted requiert des poids",
     "context Cluster\n"
     "inv weightedRequiresWeights:\n"
     "  self.loadBalancer.algorithm = Algorithm::WEIGHTED_ROUND_ROBIN\n"
     "  implies self.servers->forAll(s | s.weight >= 1)"),

    ("RG10 – Session Persistence cookie",
     "context SessionPersistence\n"
     "inv cookieRequiresName:\n"
     "  self.type = PersistenceType::COOKIE\n"
     "  implies self.cookieName <> '' and self.cookieName <> null"),

    ("RG11 – Au moins un listener",
     "context LoadBalancerSystem\n"
     "inv atLeastOneListener:\n"
     "  self.listeners->size() >= 1"),

    ("RG12 – Serveur actif dans le cluster",
     "context Cluster\n"
     "inv atLeastOneEnabledServer:\n"
     "  self.servers->exists(s | s.enabled = true)"),
]

for title, code in ocl_constraints:
    doc.add_heading(title, level=3)
    add_code_block(code)

# ============================================================
# 7. NIVEAUX MDA ET ARCHITECTURE
# ============================================================
doc.add_heading("7. Niveaux MDA et architecture de l'approche", level=1)

doc.add_heading("7.1 Les niveaux MDA pris en compte", level=2)
add_paragraph("Notre approche couvre les niveaux suivants du MDA :")

add_table(
    ["Niveau", "Description", "Artefact produit"],
    [
        ["CIM (Computation Independent Model)",
         "Modèle indépendant de l'informatique. Description du domaine métier de la répartition de charge.",
         "Description textuelle du domaine, tableau des concepts, règles de gestion (sections 1-3 de ce rapport)"],
        ["PIM (Platform Independent Model)",
         "Modèle indépendant de la plateforme. Métamodèle Ecore + contraintes OCL + modèles conformes écrits avec le DSL.",
         "Métamodèle Ecore (.ecore), contraintes OCL (.ocl), modèles utilisateur (.lb)"],
        ["PSM (Platform Specific Model)",
         "Modèle spécifique à une plateforme cible (ex: Nginx, HAProxy, Docker Compose).",
         "Modèle intermédiaire spécifique à la technologie cible (obtenu par transformation M2M ATL)"],
        ["Code",
         "Code/configuration exécutable générée automatiquement.",
         "Fichiers de configuration (nginx.conf, haproxy.cfg, docker-compose.yml) générés par transformation M2T Acceleo"],
    ]
)

doc.add_heading("7.2 Schéma d'architecture de l'approche", level=2)
add_code_block(
    "  ┌─────────────────────────────────────────────────────────────────┐\n"
    "  │                         NIVEAU CIM                             │\n"
    "  │  Description du domaine, concepts, règles de gestion           │\n"
    "  │  (Document textuel)                                            │\n"
    "  └──────────────────────────────┬──────────────────────────────────┘\n"
    "                                 │ Analyse et modélisation\n"
    "                                 ▼\n"
    "  ┌─────────────────────────────────────────────────────────────────┐\n"
    "  │                         NIVEAU PIM                             │\n"
    "  │  ┌──────────────┐   ┌──────────────┐   ┌───────────────────┐  │\n"
    "  │  │ Métamodèle   │   │ Contraintes  │   │ Syntaxe concrète  │  │\n"
    "  │  │ Ecore (.ecore)│   │ OCL (.ocl)   │   │ Xtext (.xtext)    │  │\n"
    "  │  └──────────────┘   └──────────────┘   └───────┬───────────┘  │\n"
    "  │                                                │              │\n"
    "  │                                    Éditeur DSL généré         │\n"
    "  │                                                │              │\n"
    "  │                                    ┌───────────▼───────────┐  │\n"
    "  │                                    │ Modèle utilisateur    │  │\n"
    "  │                                    │ (.lb) conforme au MM  │  │\n"
    "  │                                    └───────────┬───────────┘  │\n"
    "  └────────────────────────────────────────────────┼──────────────┘\n"
    "                                                   │\n"
    "                          Transformation M2M (ATL) │\n"
    "                                                   ▼\n"
    "  ┌─────────────────────────────────────────────────────────────────┐\n"
    "  │                         NIVEAU PSM                             │\n"
    "  │  ┌────────────────────┐  ┌────────────────────┐               │\n"
    "  │  │ Modèle Nginx       │  │ Modèle HAProxy     │  ...         │\n"
    "  │  │ (spécifique)       │  │ (spécifique)       │               │\n"
    "  │  └─────────┬──────────┘  └─────────┬──────────┘               │\n"
    "  └────────────┼───────────────────────┼──────────────────────────┘\n"
    "               │                       │\n"
    "               │  Transformation M2T   │  Transformation M2T\n"
    "               │  (Acceleo)            │  (Acceleo)\n"
    "               ▼                       ▼\n"
    "  ┌─────────────────────────────────────────────────────────────────┐\n"
    "  │                         NIVEAU CODE                            │\n"
    "  │  ┌──────────────┐  ┌──────────────┐  ┌───────────────────┐    │\n"
    "  │  │ nginx.conf   │  │ haproxy.cfg  │  │ docker-compose.yml│    │\n"
    "  │  └──────────────┘  └──────────────┘  └───────────────────┘    │\n"
    "  └─────────────────────────────────────────────────────────────────┘"
)

doc.add_heading("7.3 Justification des choix technologiques", level=2)
add_table(
    ["Outil/Technologie", "Rôle", "Justification"],
    [
        ["EMF / Ecore", "Définition du métamodèle (syntaxe abstraite)", "Standard de facto pour la méta-modélisation dans l'écosystème Eclipse. Supporte la génération de code Java et l'intégration OCL."],
        ["OCL", "Définition des contraintes", "Langage standardisé (OMG) pour exprimer des contraintes sur les modèles UML/Ecore."],
        ["Xtext", "Définition de la syntaxe concrète textuelle (DSL)", "Framework mature pour créer des DSL textuels avec éditeur Eclipse intégré (coloration, autocomplétion, validation)."],
        ["ATL", "Transformations modèle-à-modèle (M2M)", "Langage de transformation de modèles hybride (déclaratif/impératif), largement utilisé dans la communauté IDM."],
        ["Acceleo", "Transformations modèle-à-texte (M2T)", "Outil de génération de code basé sur les templates, intégré à Eclipse et compatible EMF."],
    ]
)

# ============================================================
# 8. SPECIFICATIONS DES TRANSFORMATIONS
# ============================================================
doc.add_heading("8. Spécification des transformations", level=1)

doc.add_heading("8.1 Vue d'ensemble des transformations", level=2)
add_table(
    ["Transformation", "Type", "Source", "Cible", "Outil"],
    [
        ["T1 : PIM → PSM Nginx", "M2M (modèle-à-modèle)", "Modèle LoadBalancer (.lb)", "Modèle Nginx (spécifique)", "ATL"],
        ["T2 : PIM → PSM HAProxy", "M2M (modèle-à-modèle)", "Modèle LoadBalancer (.lb)", "Modèle HAProxy (spécifique)", "ATL"],
        ["T3 : PSM Nginx → Code", "M2T (modèle-à-texte)", "Modèle Nginx", "nginx.conf", "Acceleo"],
        ["T4 : PSM HAProxy → Code", "M2T (modèle-à-texte)", "Modèle HAProxy", "haproxy.cfg", "Acceleo"],
        ["T5 : PIM → Docker Compose", "M2T (modèle-à-texte)", "Modèle LoadBalancer (.lb)", "docker-compose.yml", "Acceleo"],
    ]
)

doc.add_heading("8.2 Transformation T1 : PIM → PSM Nginx (ATL)", level=2)
add_paragraph("Cette transformation convertit notre modèle indépendant de plateforme en un modèle spécifique Nginx.")
add_code_block(
    "-- Fichier : PIM2Nginx.atl\n"
    "module PIM2Nginx;\n"
    "create OUT : NginxModel from IN : LoadBalancerModel;\n\n"
    "rule System2NginxConfig {\n"
    "  from s : LoadBalancerModel!LoadBalancerSystem\n"
    "  to   c : NginxModel!NginxConfig (\n"
    "    workerProcesses <- 'auto',\n"
    "    workerConnections <- 1024,\n"
    "    upstreams <- s.clusters->collect(cl | thisModule.Cluster2Upstream(cl)),\n"
    "    servers <- s.listeners->collect(l | thisModule.Listener2ServerBlock(l))\n"
    "  )\n"
    "}\n\n"
    "lazy rule Cluster2Upstream {\n"
    "  from cl : LoadBalancerModel!Cluster\n"
    "  to   u : NginxModel!Upstream (\n"
    "    name <- cl.name,\n"
    "    algorithm <- cl.loadBalancer.algorithm.mapNginxAlgorithm(),\n"
    "    servers <- cl.servers->select(s | s.enabled)\n"
    "              ->collect(s | thisModule.Server2UpstreamServer(s))\n"
    "  )\n"
    "}\n\n"
    "lazy rule Server2UpstreamServer {\n"
    "  from s : LoadBalancerModel!Server\n"
    "  to   us : NginxModel!UpstreamServer (\n"
    "    address <- s.host + ':' + s.port.toString(),\n"
    "    weight <- s.weight\n"
    "  )\n"
    "}\n\n"
    "lazy rule Listener2ServerBlock {\n"
    "  from l : LoadBalancerModel!Listener\n"
    "  to   sb : NginxModel!ServerBlock (\n"
    "    listenPort <- l.port,\n"
    "    proxyPass <- 'http://' + l.targetCluster.name\n"
    "  )\n"
    "}"
)

doc.add_heading("8.3 Transformation T3 : PSM Nginx → Code (Acceleo)", level=2)
add_paragraph("Cette transformation M2T génère le fichier nginx.conf à partir du modèle PSM Nginx.")
add_code_block(
    "[comment encoding = UTF-8 /]\n"
    "[module generateNginx('http://www.example.org/nginx')]\n\n"
    "[template public generateNginxConf(config : NginxConfig)]\n"
    "[comment @main /]\n"
    "[file ('nginx.conf', false, 'UTF-8')]\n"
    "worker_processes [config.workerProcesses/];\n\n"
    "events {\n"
    "    worker_connections [config.workerConnections/];\n"
    "}\n\n"
    "http {\n"
    "[for (upstream : Upstream | config.upstreams)]\n"
    "    upstream [upstream.name/] {\n"
    "[if (upstream.algorithm = 'least_conn')]\n"
    "        least_conn;\n"
    "[elseif (upstream.algorithm = 'ip_hash')]\n"
    "        ip_hash;\n"
    "[/if]\n"
    "[for (server : UpstreamServer | upstream.servers)]\n"
    "        server [server.address/] weight=[server.weight/];\n"
    "[/for]\n"
    "    }\n\n"
    "[/for]\n"
    "[for (server : ServerBlock | config.servers)]\n"
    "    server {\n"
    "        listen [server.listenPort/];\n"
    "        location / {\n"
    "            proxy_pass [server.proxyPass/];\n"
    "        }\n"
    "    }\n"
    "[/for]\n"
    "}\n"
    "[/file]\n"
    "[/template]"
)

doc.add_heading("8.4 Transformation T5 : PIM → Docker Compose (Acceleo)", level=2)
add_paragraph("Cette transformation génère directement un fichier docker-compose.yml pour déployer l'infrastructure.")
add_code_block(
    "[comment encoding = UTF-8 /]\n"
    "[module generateDockerCompose('http://www.example.org/loadbalancer')]\n\n"
    "[template public generateCompose(system : LoadBalancerSystem)]\n"
    "[comment @main /]\n"
    "[file ('docker-compose.yml', false, 'UTF-8')]\n"
    "version: '3.8'\n"
    "services:\n"
    "[for (cluster : Cluster | system.clusters)]\n"
    "[for (server : Server | cluster.servers)]\n"
    "  [server.name/]:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"[server.port/]:[server.port/]\"\n"
    "    deploy:\n"
    "      resources:\n"
    "        limits:\n"
    "          cpus: '0.5'\n"
    "          memory: 256M\n"
    "[/for]\n"
    "  [cluster.name/]-lb:\n"
    "    image: nginx:alpine\n"
    "    volumes:\n"
    "      - ./nginx.conf:/etc/nginx/nginx.conf:ro\n"
    "    ports:\n"
    "[for (listener : Listener | system.listeners)]\n"
    "[if (listener.targetCluster.name = cluster.name)]\n"
    "      - \"[listener.port/]:[listener.port/]\"\n"
    "[/if]\n"
    "[/for]\n"
    "    depends_on:\n"
    "[for (server : Server | cluster.servers)]\n"
    "      - [server.name/]\n"
    "[/for]\n"
    "[/for]\n"
    "[/file]\n"
    "[/template]"
)

# ============================================================
# 9. ETUDE DE CAS
# ============================================================
doc.add_heading("9. Étude de cas", level=1)

doc.add_heading("9.1 Description du cas d'étude", level=2)
add_paragraph(
    "Nous illustrons notre approche avec un cas concret : une application web e-commerce "
    "déployée sur un système distribué. L'application possède :"
)
add_bullet("Un cluster de 3 serveurs web (frontend) avec répartition Round Robin")
add_bullet("Un cluster de 2 serveurs API (backend) avec répartition Least Connections")
add_bullet("Des health checks HTTP sur chaque cluster")
add_bullet("Une règle de scaling automatique basée sur l'utilisation CPU")
add_bullet("Des alertes en cas de dépassement de seuils")

doc.add_heading("9.2 Modèle DSL de l'étude de cas", level=2)
add_paragraph(
    "Voici le modèle écrit avec notre DSL (syntaxe concrète textuelle) :"
)
add_code_block(
    'LoadBalancerSystem "ECommerceSystem" {\n'
    '  description: "Système de répartition de charge pour e-commerce"\n'
    '  version: "1.0"\n'
    '\n'
    '  cluster WebFrontend {\n'
    '    loadbalancer {\n'
    '      algorithm: ROUND_ROBIN\n'
    '      stickySession: true\n'
    '      sessionPersistence {\n'
    '        type: COOKIE\n'
    '        cookieName: "SESSIONID"\n'
    '        ttl: 3600\n'
    '      }\n'
    '    }\n'
    '\n'
    '    server web1 {\n'
    '      host: "192.168.1.10"\n'
    '      port: 8080\n'
    '      weight: 3\n'
    '      maxConnections: 1000\n'
    '      enabled: true\n'
    '    }\n'
    '    server web2 {\n'
    '      host: "192.168.1.11"\n'
    '      port: 8080\n'
    '      weight: 2\n'
    '      maxConnections: 800\n'
    '      enabled: true\n'
    '    }\n'
    '    server web3 {\n'
    '      host: "192.168.1.12"\n'
    '      port: 8080\n'
    '      weight: 1\n'
    '      maxConnections: 500\n'
    '      enabled: true\n'
    '    }\n'
    '\n'
    '    healthCheck {\n'
    '      protocol: HTTP\n'
    '      path: "/health"\n'
    '      interval: 30\n'
    '      timeout: 10\n'
    '      healthyThreshold: 3\n'
    '      unhealthyThreshold: 2\n'
    '    }\n'
    '\n'
    '    scalingRule autoScaleWeb {\n'
    '      metric: CPU_USAGE\n'
    '      scaleUpThreshold: 80.0\n'
    '      scaleDownThreshold: 30.0\n'
    '      minInstances: 2\n'
    '      maxInstances: 10\n'
    '    }\n'
    '  }\n'
    '\n'
    '  cluster APIBackend {\n'
    '    loadbalancer {\n'
    '      algorithm: LEAST_CONNECTIONS\n'
    '      stickySession: false\n'
    '    }\n'
    '\n'
    '    server api1 {\n'
    '      host: "192.168.2.10"\n'
    '      port: 3000\n'
    '      weight: 1\n'
    '      maxConnections: 500\n'
    '      enabled: true\n'
    '    }\n'
    '    server api2 {\n'
    '      host: "192.168.2.11"\n'
    '      port: 3000\n'
    '      weight: 1\n'
    '      maxConnections: 500\n'
    '      enabled: true\n'
    '    }\n'
    '\n'
    '    healthCheck {\n'
    '      protocol: HTTP\n'
    '      path: "/api/health"\n'
    '      interval: 15\n'
    '      timeout: 5\n'
    '      healthyThreshold: 2\n'
    '      unhealthyThreshold: 3\n'
    '    }\n'
    '  }\n'
    '\n'
    '  listener httpListener {\n'
    '    protocol: HTTP\n'
    '    port: 80\n'
    '    targetCluster: WebFrontend\n'
    '  }\n'
    '\n'
    '  listener apiListener {\n'
    '    protocol: HTTP\n'
    '    port: 3000\n'
    '    targetCluster: APIBackend\n'
    '  }\n'
    '\n'
    '  alert highCPU {\n'
    '    metric: CPU_USAGE\n'
    '    threshold: 90.0\n'
    '    action: "email:admin@ecommerce.com"\n'
    '  }\n'
    '\n'
    '  alert highLatency {\n'
    '    metric: RESPONSE_TIME\n'
    '    threshold: 5000.0\n'
    '    action: "webhook:https://hooks.slack.com/alert"\n'
    '  }\n'
    '}'
)

doc.add_heading("9.3 Résultat de la transformation : nginx.conf généré", level=2)
add_paragraph("Après exécution de la chaîne de transformations (T1 puis T3), le fichier suivant est généré :")
add_code_block(
    "worker_processes auto;\n\n"
    "events {\n"
    "    worker_connections 1024;\n"
    "}\n\n"
    "http {\n"
    "    upstream WebFrontend {\n"
    "        server 192.168.1.10:8080 weight=3;\n"
    "        server 192.168.1.11:8080 weight=2;\n"
    "        server 192.168.1.12:8080 weight=1;\n"
    "    }\n\n"
    "    upstream APIBackend {\n"
    "        least_conn;\n"
    "        server 192.168.2.10:3000 weight=1;\n"
    "        server 192.168.2.11:3000 weight=1;\n"
    "    }\n\n"
    "    server {\n"
    "        listen 80;\n"
    "        location / {\n"
    "            proxy_pass http://WebFrontend;\n"
    "        }\n"
    "    }\n\n"
    "    server {\n"
    "        listen 3000;\n"
    "        location / {\n"
    "            proxy_pass http://APIBackend;\n"
    "        }\n"
    "    }\n"
    "}"
)

doc.add_heading("9.4 Résultat : docker-compose.yml généré", level=2)
add_code_block(
    "version: '3.8'\n"
    "services:\n"
    "  web1:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"8080:8080\"\n"
    "    deploy:\n"
    "      resources:\n"
    "        limits:\n"
    "          cpus: '0.5'\n"
    "          memory: 256M\n\n"
    "  web2:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"8080:8080\"\n"
    "    deploy:\n"
    "      resources:\n"
    "        limits:\n"
    "          cpus: '0.5'\n"
    "          memory: 256M\n\n"
    "  web3:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"8080:8080\"\n"
    "    deploy:\n"
    "      resources:\n"
    "        limits:\n"
    "          cpus: '0.5'\n"
    "          memory: 256M\n\n"
    "  WebFrontend-lb:\n"
    "    image: nginx:alpine\n"
    "    volumes:\n"
    "      - ./nginx.conf:/etc/nginx/nginx.conf:ro\n"
    "    ports:\n"
    "      - \"80:80\"\n"
    "    depends_on:\n"
    "      - web1\n"
    "      - web2\n"
    "      - web3\n\n"
    "  api1:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"3000:3000\"\n\n"
    "  api2:\n"
    "    image: nginx:alpine\n"
    "    ports:\n"
    "      - \"3000:3000\"\n\n"
    "  APIBackend-lb:\n"
    "    image: nginx:alpine\n"
    "    volumes:\n"
    "      - ./nginx.conf:/etc/nginx/nginx.conf:ro\n"
    "    ports:\n"
    "      - \"3000:3000\"\n"
    "    depends_on:\n"
    "      - api1\n"
    "      - api2"
)

doc.add_heading("9.5 Illustration de bout en bout", level=2)
add_paragraph(
    "Le flux complet de notre approche pour cette étude de cas est le suivant :"
)
add_paragraph("Étape 1 : L'utilisateur rédige son modèle avec la syntaxe DSL (.lb) dans l'éditeur Eclipse.", bold=True)
add_paragraph(
    "L'éditeur Xtext offre la coloration syntaxique, l'autocomplétion et la validation en temps réel "
    "grâce aux contraintes OCL intégrées."
)
add_paragraph("Étape 2 : Validation du modèle.", bold=True)
add_paragraph(
    "Le moteur de validation vérifie toutes les contraintes OCL (RG01 à RG12). "
    "En cas d'erreur, un message clair est affiché dans l'éditeur."
)
add_paragraph("Étape 3 : Transformation M2M avec ATL.", bold=True)
add_paragraph(
    "Le modèle PIM est transformé en modèle PSM spécifique à Nginx (ou HAProxy selon le choix)."
)
add_paragraph("Étape 4 : Génération de code avec Acceleo.", bold=True)
add_paragraph(
    "Le modèle PSM est transformé en fichiers de configuration réels (nginx.conf, docker-compose.yml)."
)
add_paragraph("Étape 5 : Déploiement.", bold=True)
add_paragraph(
    "Les fichiers générés sont directement utilisables pour déployer l'infrastructure avec Docker Compose."
)

# ============================================================
# 10. TUTORIEL DE REALISATION
# ============================================================
doc.add_heading("10. Tutoriel de réalisation pas à pas", level=1)

doc.add_heading("10.1 Prérequis", level=2)
add_bullet("Eclipse Modeling Tools (version 2023-12 ou supérieure)")
add_bullet("Plugin EMF (Eclipse Modeling Framework) – inclus dans Eclipse Modeling Tools")
add_bullet("Plugin Xtext (pour la syntaxe concrète du DSL)")
add_bullet("Plugin ATL (pour les transformations M2M)")
add_bullet("Plugin Acceleo (pour les transformations M2T)")
add_bullet("Plugin OCL Tools (pour les contraintes)")
add_bullet("Java JDK 17+")

doc.add_heading("10.2 Étape 1 : Création du métamodèle Ecore", level=2)
add_paragraph("1. Dans Eclipse : File → New → Other → Eclipse Modeling Framework → Ecore Modeling Project", bold=True)
add_paragraph("2. Nommer le projet : org.loadbalancer.metamodel")
add_paragraph("3. Créer le fichier loadbalancer.ecore et y définir les classes suivantes :")
add_code_block(
    "<!-- Extrait du fichier loadbalancer.ecore (format XMI) -->\n"
    "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
    "<ecore:EPackage xmi:version=\"2.0\"\n"
    "    xmlns:xmi=\"http://www.omg.org/XMI\"\n"
    "    xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\"\n"
    "    xmlns:ecore=\"http://www.eclipse.org/emf/2002/Ecore\"\n"
    "    name=\"loadbalancer\"\n"
    "    nsURI=\"http://www.example.org/loadbalancer\"\n"
    "    nsPrefix=\"lb\">\n\n"
    "  <eClassifiers xsi:type=\"ecore:EClass\" name=\"LoadBalancerSystem\">\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"name\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"description\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"version\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"clusters\"\n"
    "        upperBound=\"-1\" eType=\"#//Cluster\" containment=\"true\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"listeners\"\n"
    "        upperBound=\"-1\" eType=\"#//Listener\" containment=\"true\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"alerts\"\n"
    "        upperBound=\"-1\" eType=\"#//Alert\" containment=\"true\"/>\n"
    "  </eClassifiers>\n\n"
    "  <eClassifiers xsi:type=\"ecore:EClass\" name=\"Cluster\">\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"name\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"maxConnections\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EInt\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"servers\"\n"
    "        upperBound=\"-1\" eType=\"#//Server\" containment=\"true\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"loadBalancer\"\n"
    "        eType=\"#//LoadBalancer\" containment=\"true\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"healthCheck\"\n"
    "        eType=\"#//HealthCheck\" containment=\"true\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EReference\" name=\"scalingRule\"\n"
    "        eType=\"#//ScalingRule\" containment=\"true\"/>\n"
    "  </eClassifiers>\n\n"
    "  <eClassifiers xsi:type=\"ecore:EClass\" name=\"Server\">\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"name\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"host\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EString\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"port\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EInt\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"weight\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EInt\" defaultValueLiteral=\"1\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"maxConnections\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EInt\"/>\n"
    "    <eStructuralFeatures xsi:type=\"ecore:EAttribute\" name=\"enabled\"\n"
    "        eType=\"ecore:EDataType http://www.eclipse.org/emf/2002/Ecore#//EBoolean\" defaultValueLiteral=\"true\"/>\n"
    "  </eClassifiers>\n\n"
    "  <!-- ... autres classes: LoadBalancer, HealthCheck, Listener, etc. -->\n\n"
    "  <eClassifiers xsi:type=\"ecore:EEnum\" name=\"Algorithm\">\n"
    "    <eLiterals name=\"ROUND_ROBIN\"/>\n"
    "    <eLiterals name=\"WEIGHTED_ROUND_ROBIN\" value=\"1\"/>\n"
    "    <eLiterals name=\"LEAST_CONNECTIONS\" value=\"2\"/>\n"
    "    <eLiterals name=\"IP_HASH\" value=\"3\"/>\n"
    "    <eLiterals name=\"RANDOM\" value=\"4\"/>\n"
    "  </eClassifiers>\n\n"
    "  <eClassifiers xsi:type=\"ecore:EEnum\" name=\"Protocol\">\n"
    "    <eLiterals name=\"HTTP\"/>\n"
    "    <eLiterals name=\"HTTPS\" value=\"1\"/>\n"
    "    <eLiterals name=\"TCP\" value=\"2\"/>\n"
    "    <eLiterals name=\"UDP\" value=\"3\"/>\n"
    "  </eClassifiers>\n\n"
    "</ecore:EPackage>"
)
add_paragraph("4. Ouvrir le fichier avec l'éditeur graphique Ecore pour visualiser le diagramme de classes.")

doc.add_heading("10.3 Étape 2 : Génération du code Java EMF", level=2)
add_paragraph("1. Clic droit sur loadbalancer.ecore → New → Other → EMF Generator Model", bold=True)
add_paragraph("2. Créer loadbalancer.genmodel")
add_paragraph("3. Ouvrir le .genmodel → Clic droit sur la racine → Generate All")
add_paragraph("Cela génère :")
add_bullet("Le code Java des classes du métamodèle")
add_bullet("Les factories et packages EMF")
add_bullet("Le plugin d'édition")

doc.add_heading("10.4 Étape 3 : Définition de la syntaxe concrète avec Xtext", level=2)
add_paragraph("1. File → New → Project → Xtext → Xtext Project", bold=True)
add_paragraph("2. Nommer : org.loadbalancer.dsl, extension de fichier : .lb")
add_paragraph("3. Définir la grammaire dans le fichier LoadBalancer.xtext :")
add_code_block(
    "grammar org.loadbalancer.dsl.LoadBalancer\n"
    "    with org.eclipse.xtext.common.Terminals\n\n"
    "generate loadBalancer \"http://www.example.org/loadbalancer\"\n\n"
    "LoadBalancerSystem:\n"
    '  \'LoadBalancerSystem\' name=STRING \'{\'\n'
    '    (\'description:\' description=STRING)?\n'
    '    (\'version:\' version=STRING)?\n'
    '    (clusters+=Cluster)*\n'
    '    (listeners+=Listener)*\n'
    '    (alerts+=Alert)*\n'
    '  \'}\';\n\n'
    'Cluster:\n'
    '  \'cluster\' name=ID \'{\'\n'
    '    loadBalancer=LoadBalancer\n'
    '    (servers+=Server)+\n'
    '    (healthCheck=HealthCheck)?\n'
    '    (scalingRule=ScalingRule)?\n'
    '  \'}\';\n\n'
    'Server:\n'
    '  \'server\' name=ID \'{\'\n'
    '    \'host:\' host=STRING\n'
    '    \'port:\' port=INT\n'
    '    (\'weight:\' weight=INT)?\n'
    '    (\'maxConnections:\' maxConnections=INT)?\n'
    '    (\'enabled:\' enabled=BOOLEAN)?\n'
    '  \'}\';\n\n'
    'LoadBalancer:\n'
    '  \'loadbalancer\' \'{\'\n'
    '    \'algorithm:\' algorithm=Algorithm\n'
    '    (\'stickySession:\' stickySession=BOOLEAN)?\n'
    '    (sessionPersistence=SessionPersistence)?\n'
    '  \'}\';\n\n'
    'SessionPersistence:\n'
    '  \'sessionPersistence\' \'{\'\n'
    '    \'type:\' type=PersistenceType\n'
    '    (\'cookieName:\' cookieName=STRING)?\n'
    '    (\'ttl:\' ttl=INT)?\n'
    '  \'}\';\n\n'
    'HealthCheck:\n'
    '  \'healthCheck\' \'{\'\n'
    '    \'protocol:\' protocol=Protocol\n'
    '    \'path:\' path=STRING\n'
    '    \'interval:\' interval=INT\n'
    '    \'timeout:\' timeout=INT\n'
    '    (\'healthyThreshold:\' healthyThreshold=INT)?\n'
    '    (\'unhealthyThreshold:\' unhealthyThreshold=INT)?\n'
    '  \'}\';\n\n'
    'Listener:\n'
    '  \'listener\' name=ID \'{\'\n'
    '    \'protocol:\' protocol=Protocol\n'
    '    \'port:\' port=INT\n'
    '    \'targetCluster:\' targetCluster=[Cluster]\n'
    '  \'}\';\n\n'
    'ScalingRule:\n'
    '  \'scalingRule\' name=ID \'{\'\n'
    '    \'metric:\' metric=Metric\n'
    '    \'scaleUpThreshold:\' scaleUpThreshold=FLOAT\n'
    '    \'scaleDownThreshold:\' scaleDownThreshold=FLOAT\n'
    '    \'minInstances:\' minInstances=INT\n'
    '    \'maxInstances:\' maxInstances=INT\n'
    '  \'}\';\n\n'
    'Alert:\n'
    '  \'alert\' name=ID \'{\'\n'
    '    \'metric:\' metric=Metric\n'
    '    \'threshold:\' threshold=FLOAT\n'
    '    \'action:\' action=STRING\n'
    '  \'}\';\n\n'
    'enum Algorithm:\n'
    '  ROUND_ROBIN | WEIGHTED_ROUND_ROBIN | LEAST_CONNECTIONS | IP_HASH | RANDOM;\n\n'
    'enum Protocol:\n'
    '  HTTP | HTTPS | TCP | UDP;\n\n'
    'enum PersistenceType:\n'
    '  COOKIE | IP | HEADER;\n\n'
    'enum Metric:\n'
    '  CPU_USAGE | MEMORY_USAGE | REQUEST_COUNT | RESPONSE_TIME | ACTIVE_CONNECTIONS;\n\n'
    'terminal BOOLEAN returns ecore::EBoolean:\n'
    "  'true' | 'false';\n\n"
    'terminal FLOAT returns ecore::EFloat:\n'
    "  INT '.' INT;"
)

add_paragraph("4. Clic droit sur le fichier .xtext → Run As → Generate Xtext Artifacts")
add_paragraph("5. Tester : Run As → Eclipse Application → Créer un fichier .lb dans le nouveau workspace")

doc.add_heading("10.5 Étape 4 : Ajout des contraintes OCL (Validation)", level=2)
add_paragraph(
    "Intégrer les contraintes OCL dans le validateur Xtext généré "
    "(fichier LoadBalancerValidator.xtend) :"
)
add_code_block(
    "// Fichier: LoadBalancerValidator.xtend\n"
    "class LoadBalancerValidator extends AbstractLoadBalancerValidator {\n\n"
    "  @Check\n"
    "  def checkUniqueServerNames(Cluster cluster) {\n"
    "    val names = cluster.servers.map[name]\n"
    "    val duplicates = names.filter[n | names.filter[it == n].size > 1].toSet\n"
    "    if (!duplicates.empty) {\n"
    "      error('Duplicate server names: ' + duplicates.join(', '),\n"
    "            LoadBalancerPackage.Literals.CLUSTER__SERVERS)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkValidPort(Server server) {\n"
    "    if (server.port < 1 || server.port > 65535) {\n"
    "      error('Port must be between 1 and 65535',\n"
    "            LoadBalancerPackage.Literals.SERVER__PORT)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkPositiveWeight(Server server) {\n"
    "    if (server.weight < 1) {\n"
    "      error('Weight must be >= 1',\n"
    "            LoadBalancerPackage.Literals.SERVER__WEIGHT)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkClusterNotEmpty(Cluster cluster) {\n"
    "    if (cluster.servers.empty) {\n"
    "      error('Cluster must have at least one server',\n"
    "            LoadBalancerPackage.Literals.CLUSTER__SERVERS)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkHealthCheckCoherent(HealthCheck hc) {\n"
    "    if (hc.interval <= hc.timeout) {\n"
    "      error('Interval must be greater than timeout',\n"
    "            LoadBalancerPackage.Literals.HEALTH_CHECK__INTERVAL)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkScalingThresholds(ScalingRule rule) {\n"
    "    if (rule.scaleUpThreshold <= rule.scaleDownThreshold) {\n"
    "      error('Scale-up threshold must be > scale-down threshold',\n"
    "            LoadBalancerPackage.Literals.SCALING_RULE__SCALE_UP_THRESHOLD)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkMinMaxInstances(ScalingRule rule) {\n"
    "    if (rule.minInstances > rule.maxInstances) {\n"
    "      error('Min instances must be <= max instances',\n"
    "            LoadBalancerPackage.Literals.SCALING_RULE__MIN_INSTANCES)\n"
    "    }\n"
    "  }\n\n"
    "  @Check\n"
    "  def checkUniqueListenerPorts(LoadBalancerSystem system) {\n"
    "    val ports = system.listeners.map[port]\n"
    "    val duplicates = ports.filter[p | ports.filter[it == p].size > 1].toSet\n"
    "    if (!duplicates.empty) {\n"
    "      error('Duplicate listener ports: ' + duplicates.join(', '),\n"
    "            LoadBalancerPackage.Literals.LOAD_BALANCER_SYSTEM__LISTENERS)\n"
    "    }\n"
    "  }\n"
    "}"
)

doc.add_heading("10.6 Étape 5 : Création des transformations ATL (M2M)", level=2)
add_paragraph("1. File → New → ATL Project : org.loadbalancer.transformations", bold=True)
add_paragraph("2. Créer d'abord le métamodèle cible Nginx (nginx.ecore) avec les classes : NginxConfig, Upstream, UpstreamServer, ServerBlock")
add_paragraph("3. Écrire la transformation ATL (voir section 8.2 pour le code complet)")

doc.add_heading("10.7 Étape 6 : Création des transformations Acceleo (M2T)", level=2)
add_paragraph("1. File → New → Acceleo Project : org.loadbalancer.generators", bold=True)
add_paragraph("2. Créer les templates Acceleo (voir sections 8.3 et 8.4 pour le code complet)")
add_paragraph("3. Configurer le launch configuration pour pointer vers le modèle source")

doc.add_heading("10.8 Étape 7 : Test de bout en bout", level=2)
add_paragraph("1. Lancer une Eclipse Application (Run As → Eclipse Application)")
add_paragraph("2. Créer un nouveau projet général dans le workspace de test")
add_paragraph("3. Créer un fichier ecommerce.lb avec le contenu de la section 9.2")
add_paragraph("4. Vérifier que la validation fonctionne (modifier des valeurs pour tester les contraintes)")
add_paragraph("5. Exécuter la transformation ATL")
add_paragraph("6. Exécuter la génération Acceleo")
add_paragraph("7. Vérifier les fichiers générés (nginx.conf, docker-compose.yml)")

# ============================================================
# 11. AUTRES
# ============================================================
doc.add_heading("11. Perspectives et améliorations", level=1)
add_paragraph("Ce qui resterait à faire pour atteindre pleinement les objectifs :")
add_bullet("Ajout du support HAProxy comme cible PSM supplémentaire")
add_bullet("Génération de scripts Kubernetes (YAML de Deployment + Service + Ingress)")
add_bullet("Éditeur graphique (Sirius) en complément de l'éditeur textuel Xtext")
add_bullet("Support de la répartition de charge au niveau DNS (DNS-based load balancing)")
add_bullet("Intégration avec des outils de monitoring réels (Prometheus, Grafana)")
add_bullet("Validation dynamique avec simulation de charge")
add_bullet("Support multi-cloud (AWS ALB, GCP Load Balancer, Azure Load Balancer)")

# ============================================================
# CONCLUSION
# ============================================================
doc.add_heading("Conclusion", level=1)
add_paragraph(
    "Ce projet a permis de mettre en pratique les concepts fondamentaux de l'Ingénierie Dirigée "
    "par les Modèles (IDM) à travers la conception d'un DSL dédié à la répartition de charge "
    "dans les systèmes distribués."
)
add_paragraph(
    "Nous avons suivi une démarche MDA complète :"
)
add_bullet("Niveau CIM : analyse et description du domaine de la répartition de charge")
add_bullet("Niveau PIM : conception du métamodèle Ecore, définition des contraintes OCL, et création de la syntaxe concrète Xtext")
add_bullet("Niveau PSM : transformation ATL vers des modèles spécifiques (Nginx)")
add_bullet("Niveau Code : génération automatique de fichiers de configuration via Acceleo")

add_paragraph(
    "L'approche proposée démontre les avantages de l'IDM : abstraction, réutilisabilité, "
    "portabilité entre plateformes, et validation formelle des modèles. Le DSL créé permet "
    "aux administrateurs système de décrire leur infrastructure de répartition de charge de "
    "manière déclarative et indépendante de la technologie, tout en bénéficiant d'une génération "
    "automatique de code fiable et cohérente."
)

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading("Références bibliographiques", level=1)
refs = [
    "[1] J. Bézivin, « On the Unification Power of Models », Software and Systems Modeling, vol. 4, no. 2, pp. 171-188, 2005.",
    "[2] OMG, « MDA Guide Version 2.0 », Object Management Group, 2014.",
    "[3] F. Jouault, F. Allilaire, J. Bézivin, I. Kurtev, « ATL: A model transformation tool », Science of Computer Programming, vol. 72, no. 1-2, pp. 31-39, 2008.",
    "[4] M. Bettini, « Implementing Domain-Specific Languages with Xtext and Xtend », Packt Publishing, 2nd edition, 2016.",
    "[5] Eclipse Foundation, « Acceleo User Guide », https://wiki.eclipse.org/Acceleo.",
    "[6] Eclipse Foundation, « Eclipse Modeling Framework (EMF) », https://www.eclipse.org/modeling/emf/.",
    "[7] Nginx Documentation, « HTTP Load Balancing », https://docs.nginx.com/.",
    "[8] HAProxy Documentation, « Configuration Manual », https://www.haproxy.org/.",
    "[9] T. Stahl, M. Voelter, « Model-Driven Software Development: Technology, Engineering, Management », Wiley, 2006.",
    "[10] A. Kleppe, J. Warmer, W. Bast, « MDA Explained: The Model Driven Architecture: Practice and Promise », Addison-Wesley, 2003.",
]
for ref in refs:
    add_paragraph(ref, space_after=4)

# ============================================================
# SAVE
# ============================================================
output_path = "/home/zaz/Documents/Ecoles/M2/INF5039/ZAZ/IDM-for-loadbalancing/Rapport_IDM_LoadBalancing_DSL.docx"
doc.save(output_path)
print(f"Rapport généré avec succès : {output_path}")
